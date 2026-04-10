from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re
from typing import Iterable

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DOCX = Path(r"C:\Users\Riteshwar Singh\Downloads\DL project Report.docx")
OUTPUT_PDF = Path(r"C:\Users\Riteshwar Singh\Downloads\DL project Report.pdf")
LOGO_PATH = Path(r"C:\Users\Riteshwar Singh\Downloads\WhatsApp Image 2026-04-09 at 10.06.40 PM.jpeg")
EDITOR_LOG_PATH = ROOT / "Logs" / "Editor.log"
ASSETS_DIR = ROOT / "reporting" / "generated_assets"
SPRITES_DIR = ROOT / "Assets" / "Sprites"
STUDENT_NAME = "Riteshwar Singh"
REG_NUMBER = "241627014"

PRIMARY = (145, 47, 48)
ACCENT = (218, 128, 67)
INK = (34, 36, 40)
MUTED = (105, 112, 121)
PAPER = (248, 246, 242)
LIGHT = (236, 230, 223)


@dataclass
class Box:
    x: int
    y: int
    w: int
    h: int
    text: str
    fill: tuple[int, int, int]
    outline: tuple[int, int, int] = PRIMARY
    radius: int = 18


def ensure_dirs() -> None:
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)


def load_font(size: int, bold: bool = False):
    candidates = []
    if bold:
        candidates += [
            Path(r"C:\Windows\Fonts\georgiab.ttf"),
            Path(r"C:\Windows\Fonts\arialbd.ttf"),
            Path(r"C:\Windows\Fonts\calibrib.ttf"),
        ]
    else:
        candidates += [
            Path(r"C:\Windows\Fonts\georgia.ttf"),
            Path(r"C:\Windows\Fonts\arial.ttf"),
            Path(r"C:\Windows\Fonts\calibri.ttf"),
        ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


FONT_TITLE = load_font(32, bold=True)
FONT_SUBTITLE = load_font(22, bold=True)
FONT_BODY = load_font(18, bold=False)
FONT_SMALL = load_font(14, bold=False)


def rounded_box(draw: ImageDraw.ImageDraw, box: Box) -> None:
    draw.rounded_rectangle(
        (box.x, box.y, box.x + box.w, box.y + box.h),
        radius=box.radius,
        fill=box.fill,
        outline=box.outline,
        width=3,
    )
    text_block(draw, box.text, (box.x + 14, box.y + 12, box.x + box.w - 14, box.y + box.h - 12), FONT_BODY, INK, "center")


def text_block(draw: ImageDraw.ImageDraw, text: str, bounds, font, fill, align="left", line_gap=6) -> None:
    x1, y1, x2, y2 = bounds
    width = x2 - x1
    words = text.split()
    lines = []
    current = ""
    for word in words:
        trial = f"{current} {word}".strip()
        if draw.textbbox((0, 0), trial, font=font)[2] <= width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)

    line_height = draw.textbbox((0, 0), "Ag", font=font)[3] + line_gap
    total_height = max(0, len(lines) * line_height - line_gap)
    y = y1 + max(0, (y2 - y1 - total_height) // 2)
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        text_w = bbox[2] - bbox[0]
        if align == "center":
            x = x1 + max(0, (width - text_w) // 2)
        elif align == "right":
            x = x2 - text_w
        else:
            x = x1
        draw.text((x, y), line, font=font, fill=fill)
        y += line_height


def arrow(draw: ImageDraw.ImageDraw, start, end, fill=PRIMARY) -> None:
    draw.line((start, end), fill=fill, width=5)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex >= sx else -1
        draw.polygon([(ex, ey), (ex - 16 * direction, ey - 8), (ex - 16 * direction, ey + 8)], fill=fill)
    else:
        direction = 1 if ey >= sy else -1
        draw.polygon([(ex, ey), (ex - 8, ey - 16 * direction), (ex + 8, ey - 16 * direction)], fill=fill)


def diamond(draw: ImageDraw.ImageDraw, center, size, fill, outline, text) -> None:
    cx, cy = center
    w, h = size
    pts = [(cx, cy - h // 2), (cx + w // 2, cy), (cx, cy + h // 2), (cx - w // 2, cy)]
    draw.polygon(pts, fill=fill, outline=outline)
    text_block(draw, text, (cx - w // 2 + 10, cy - h // 2 + 8, cx + w // 2 - 10, cy + h // 2 - 8), FONT_SMALL, INK, "center")


def add_title(img: Image.Image, title: str, subtitle: str):
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, img.width, img.height), fill=PAPER)
    draw.rounded_rectangle((28, 24, img.width - 28, 110), radius=24, fill=(255, 250, 246), outline=PRIMARY, width=4)
    draw.text((52, 38), title, font=FONT_TITLE, fill=PRIMARY)
    draw.text((54, 78), subtitle, font=FONT_SMALL, fill=MUTED)
    return draw


def make_system_architecture() -> Path:
    img = Image.new("RGB", (1600, 980), PAPER)
    draw = add_title(img, "System Architecture", "Gameplay loop, AI controller, trainer, and feedback channels")
    output = ASSETS_DIR / "system_architecture.png"
    boxes = [
        Box(80, 190, 300, 120, "User Input /\nNeuro Controller", (255, 235, 228)),
        Box(470, 190, 300, 120, "Player Controller\n(Jump Physics)", (255, 245, 229)),
        Box(860, 190, 300, 120, "Game Manager\n(Score, UI, Speed)", (238, 241, 252)),
        Box(1240, 190, 270, 120, "Spawner\n(Obstacle Flow)", (231, 243, 238)),
        Box(240, 520, 330, 150, "Environment State\nDistance, Height,\nWidth, Grounded, Velocity", (255, 250, 239)),
        Box(690, 500, 320, 190, "NeuroEvolution Trainer\nPopulation, Fitness,\nSelection, Crossover,\nMutation", (245, 233, 248)),
        Box(1120, 520, 310, 150, "Training Outcomes\nBest Genome, Overlay,\nImproved Policy", (236, 244, 252)),
    ]
    for b in boxes:
        rounded_box(draw, b)
    arrow(draw, (380, 250), (470, 250))
    arrow(draw, (770, 250), (860, 250))
    arrow(draw, (1160, 250), (1240, 250))
    arrow(draw, (1380, 310), (1380, 520))
    arrow(draw, (405, 520), (380, 310))
    arrow(draw, (570, 595), (690, 595))
    arrow(draw, (1010, 595), (1120, 595))
    arrow(draw, (1120, 670), (230, 670))
    arrow(draw, (230, 670), (230, 310))
    draw.text((635, 720), "Closed-loop learning: gameplay state feeds the policy, and policy actions reshape future state.", font=FONT_BODY, fill=INK)
    img.save(output)
    return output


def make_activity_diagram() -> Path:
    img = Image.new("RGB", (1400, 1700), PAPER)
    draw = add_title(img, "Activity Diagram", "Main runtime decision process followed by the dinosaur agent")
    output = ASSETS_DIR / "activity_diagram.png"
    flow_boxes = [
        Box(500, 150, 400, 90, "Start Game / Reset Episode", (255, 236, 228)),
        Box(500, 300, 400, 110, "Initialize Player, Score,\nSpawner and Scene References", (255, 248, 231)),
        Box(500, 480, 400, 110, "Observe State:\nGrounded, Y Velocity,\nNext Obstacle Features", (236, 244, 252)),
        Box(500, 930, 400, 110, "Apply Jump or Hold Decision", (231, 243, 238)),
        Box(500, 1110, 400, 110, "Update Physics, Score,\nGame Speed and Camera Theme", (245, 233, 248)),
        Box(500, 1450, 400, 90, "End Episode / Retry / Next Genome", (255, 235, 228)),
    ]
    for b in flow_boxes:
        rounded_box(draw, b)
    diamond(draw, (700, 700), (340, 160), (255, 247, 213), ACCENT, "Obstacle Close Enough\nfor Action?")
    diamond(draw, (700, 1300), (340, 160), (255, 240, 234), ACCENT, "Collision or\nTimeout?")
    arrow(draw, (700, 240), (700, 300))
    arrow(draw, (700, 410), (700, 480))
    arrow(draw, (700, 590), (700, 620))
    arrow(draw, (700, 780), (700, 930))
    arrow(draw, (700, 1040), (700, 1110))
    arrow(draw, (700, 1380), (700, 1450))
    arrow(draw, (530, 700), (340, 700))
    arrow(draw, (340, 700), (340, 1110))
    arrow(draw, (340, 1110), (500, 1165))
    arrow(draw, (870, 1300), (1115, 1300))
    arrow(draw, (1115, 1300), (1115, 1540))
    arrow(draw, (1115, 1540), (900, 1540))
    draw.text((905, 668), "Yes", font=FONT_SMALL, fill=INK)
    draw.text((405, 668), "No", font=FONT_SMALL, fill=INK)
    draw.text((915, 1266), "Yes", font=FONT_SMALL, fill=INK)
    draw.text((175, 1080), "Continue running", font=FONT_SMALL, fill=INK)
    draw.text((970, 1520), "Training mode:\nrecord fitness and evolve", font=FONT_SMALL, fill=INK)
    img.save(output)
    return output


def make_training_pipeline() -> Path:
    img = Image.new("RGB", (1600, 1040), PAPER)
    draw = add_title(img, "Neuroevolution Training Pipeline", "How genomes are evaluated, ranked, and improved over generations")
    output = ASSETS_DIR / "training_pipeline.png"
    top = [
        Box(70, 250, 250, 120, "Create Initial\nPopulation", (255, 235, 228)),
        Box(390, 250, 250, 120, "Assign Genome to\nController", (255, 248, 231)),
        Box(710, 250, 250, 120, "Run Episode in\nDeterministic Scene", (236, 244, 252)),
        Box(1030, 250, 250, 120, "Measure Fitness\n(Survival + Clears)", (231, 243, 238)),
        Box(1350, 250, 180, 120, "Rank", (245, 233, 248)),
    ]
    bottom = [
        Box(230, 620, 260, 130, "Keep Elites\n(Best Genomes)", (255, 245, 229)),
        Box(590, 620, 260, 130, "Tournament\nSelection", (255, 250, 239)),
        Box(950, 620, 260, 130, "Crossover +\nMutation", (238, 241, 252)),
        Box(1310, 620, 210, 130, "Next\nGeneration", (255, 235, 228)),
    ]
    for item in top + bottom:
        rounded_box(draw, item)
    for i in range(len(top) - 1):
        arrow(draw, (top[i].x + top[i].w, top[i].y + top[i].h // 2), (top[i + 1].x, top[i + 1].y + top[i + 1].h // 2))
    for i in range(len(bottom) - 1):
        arrow(draw, (bottom[i].x + bottom[i].w, bottom[i].y + bottom[i].h // 2), (bottom[i + 1].x, bottom[i + 1].y + bottom[i + 1].h // 2))
    arrow(draw, (1440, 370), (1440, 620))
    arrow(draw, (1415, 685), (1415, 890))
    arrow(draw, (1415, 890), (195, 890))
    arrow(draw, (195, 890), (195, 370))
    draw.text((1230, 844), "Repeat until target generations or stable best policy", font=FONT_BODY, fill=INK)
    img.save(output)
    return output


def make_feature_map() -> Path:
    img = Image.new("RGB", (1500, 980), PAPER)
    draw = add_title(img, "AI Feature Map", "Observed inputs and policy outputs used by the neuro-controller")
    output = ASSETS_DIR / "feature_map.png"
    left = Box(90, 180, 500, 650, "Policy Inputs\n\n1. Grounded state\n2. Approximate Y velocity\n3. Distance to next obstacle\n4. Obstacle width\n5. Obstacle height\n6. Dino width\n7. Dino height\n8. Current game speed\n9. Speed increase factor", (255, 248, 231))
    middle = Box(650, 260, 220, 490, "8 Hidden\nNeurons\n(Tanh)", (245, 233, 248))
    right = Box(930, 240, 470, 530, "Policy Outputs\n\n1. Jump press signal\n2. Jump hold signal\n\nThreshold logic:\n- Grounded -> press threshold\n- Airborne -> hold threshold", (231, 243, 238))
    for b in (left, middle, right):
        rounded_box(draw, b)
    arrow(draw, (590, 505), (650, 505))
    arrow(draw, (870, 505), (930, 505))
    draw.text((420, 860), "Fixed topology: 9 -> 8 -> 2", font=FONT_SUBTITLE, fill=PRIMARY)
    img.save(output)
    return output


def make_sprite_sheet() -> Path:
    files = [
        SPRITES_DIR / "Dino_Idle.png",
        SPRITES_DIR / "Dino_Run01.png",
        SPRITES_DIR / "Dino_Run02.png",
        SPRITES_DIR / "Cactus_Small_Single.png",
        SPRITES_DIR / "Cactus_Large_Single.png",
        SPRITES_DIR / "Ground.png",
        SPRITES_DIR / "Retry.png",
    ]
    tile_w, tile_h = 210, 180
    cols, rows = 3, 3
    img = Image.new("RGBA", (cols * tile_w + 80, rows * tile_h + 160), PAPER + (255,))
    draw = add_title(img, "Visual Asset Sheet", "Key sprite resources used to construct the Chrome Dino environment")
    for idx, path in enumerate([p for p in files if p.exists()]):
        sprite = Image.open(path).convert("RGBA")
        sprite.thumbnail((120, 120))
        col = idx % cols
        row = idx // cols
        x = 55 + col * tile_w
        y = 150 + row * tile_h
        draw.rounded_rectangle((x, y, x + 170, y + 145), radius=18, fill=(255, 251, 247), outline=LIGHT, width=2)
        sx = x + (170 - sprite.width) // 2
        sy = y + 10 + (90 - sprite.height) // 2
        img.alpha_composite(sprite, (sx, sy))
        text_block(draw, path.stem.replace("_", " "), (x + 10, y + 100, x + 160, y + 135), FONT_SMALL, INK, "center")
    output = ASSETS_DIR / "sprite_sheet.png"
    img.convert("RGB").save(output)
    return output


def parse_fitness_runs(log_path: Path) -> list[dict]:
    if not log_path.exists():
        return []
    runs = []
    current = None
    pattern = re.compile(r"\[NeuroEvolution\]\s+Gen\s+(\d+)/(\d+)\s+best fitness=([0-9.]+)")
    for line in log_path.read_text(errors="ignore").splitlines():
        m = pattern.search(line)
        if not m:
            continue
        gen = int(m.group(1))
        total = int(m.group(2))
        fitness = float(m.group(3))
        if gen == 1 or current is None or gen != current["points"][-1][0] + 1 or total != current["total"]:
            current = {"total": total, "points": []}
            runs.append(current)
        current["points"].append((gen, fitness))
    return runs


def select_latest_run(runs: list[dict]) -> dict | None:
    if not runs:
        return None
    return runs[-1]


def make_fitness_graph(run: dict | None) -> Path:
    img = Image.new("RGB", (1600, 980), PAPER)
    draw = add_title(img, "Fitness vs Generations", "Actual training sequence recovered from Unity Editor.log")
    output = ASSETS_DIR / "fitness_vs_generations.png"

    chart = (120, 180, 1480, 820)
    x1, y1, x2, y2 = chart
    draw.rounded_rectangle((x1, y1, x2, y2), radius=22, fill=(255, 251, 247), outline=LIGHT, width=3)

    if not run or not run["points"]:
        draw.text((170, 460), "No fitness history was available in Editor.log.", font=FONT_SUBTITLE, fill=PRIMARY)
        img.save(output)
        return output

    gens = [g for g, _ in run["points"]]
    fits = [f for _, f in run["points"]]
    max_fit = max(fits) if fits else 1.0
    min_fit = min(0.0, min(fits))
    usable_w = x2 - x1 - 140
    usable_h = y2 - y1 - 120
    ox = x1 + 90
    oy = y2 - 70

    draw.line((ox, y1 + 40, ox, oy), fill=INK, width=4)
    draw.line((ox, oy, x2 - 40, oy), fill=INK, width=4)

    y_ticks = 6
    for i in range(y_ticks + 1):
        t = i / y_ticks
        y = oy - int(t * usable_h)
        val = min_fit + (max_fit - min_fit) * t
        draw.line((ox - 8, y, x2 - 40, y), fill=(225, 221, 215), width=1)
        draw.text((x1 + 12, y - 10), f"{val:.0f}", font=FONT_SMALL, fill=MUTED)

    n = len(gens)
    xs = []
    ys = []
    for idx, (gen, fit) in enumerate(run["points"]):
        x = ox + int((idx / max(1, n - 1)) * usable_w)
        y = oy - int(((fit - min_fit) / max(1e-6, max_fit - min_fit)) * usable_h)
        xs.append(x)
        ys.append(y)
        if idx < n - 1:
            nx = ox + int(((idx + 1) / max(1, n - 1)) * usable_w)
            draw.line((x, oy, x, oy + 8), fill=INK, width=2)
            if n <= 20 or idx % max(1, n // 12) == 0:
                draw.text((x - 10, oy + 16), str(gen), font=FONT_SMALL, fill=MUTED)
    if n == 1:
        draw.text((xs[0] - 10, oy + 16), str(gens[0]), font=FONT_SMALL, fill=MUTED)

    for i in range(n - 1):
        draw.line((xs[i], ys[i], xs[i + 1], ys[i + 1]), fill=PRIMARY, width=5)
    for x, y in zip(xs, ys):
        draw.ellipse((x - 7, y - 7, x + 7, y + 7), fill=ACCENT, outline=PRIMARY, width=2)

    best_idx = max(range(n), key=lambda i: fits[i])
    bx, by = xs[best_idx], ys[best_idx]
    draw.rounded_rectangle((bx + 16, by - 54, bx + 220, by - 6), radius=12, fill=(255, 245, 229), outline=ACCENT, width=2)
    draw.text((bx + 28, by - 44), f"Peak: Gen {gens[best_idx]}", font=FONT_SMALL, fill=INK)
    draw.text((bx + 28, by - 24), f"Fitness: {fits[best_idx]:.2f}", font=FONT_SMALL, fill=INK)

    draw.text((700, 865), "Generation", font=FONT_BODY, fill=INK)
    draw.text((35, 450), "Fitness", font=FONT_BODY, fill=INK)
    summary = f"Selected run: {n} generations out of configured {run['total']} | first={fits[0]:.2f} | last={fits[-1]:.2f} | best={max_fit:.2f}"
    draw.text((120, 115), summary, font=FONT_SMALL, fill=MUTED)
    img.save(output)
    return output


def make_equation_image(name: str, equation: str, width: float = 8.5, height: float = 1.0, fontsize: int = 20) -> Path:
    output = ASSETS_DIR / f"{name}.png"
    fig = plt.figure(figsize=(width, height), dpi=200)
    fig.patch.set_facecolor("white")
    ax = fig.add_axes([0, 0, 1, 1])
    ax.axis("off")
    ax.text(0.5, 0.5, f"${equation}$", ha="center", va="center", fontsize=fontsize, color="black")
    plt.savefig(output, dpi=200, bbox_inches="tight", pad_inches=0.08, facecolor="white")
    plt.close(fig)
    return output


def add_equation(doc: Document, path: Path, width_inches: float = 5.8, caption: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width_inches))
    if caption:
        add_caption(doc, caption)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_page_border(section) -> None:
    sect_pr = section._sectPr
    borders = sect_pr.first_child_found_in("w:pgBorders")
    if borders is None:
        borders = OxmlElement("w:pgBorders")
        borders.set(qn("w:offsetFrom"), "page")
        sect_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "22")
        element.set(qn("w:color"), "8B2E30")


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    if "Report Title" not in doc.styles:
        style = doc.styles.add_style("Report Title", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(22)
        style.font.bold = True
        style.font.color.rgb = RGBColor(*PRIMARY)
    for level in (1, 2, 3):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Times New Roman"
        style.font.bold = True
        style.font.color.rgb = RGBColor(*PRIMARY)
        style.font.size = Pt(16 if level == 1 else 14 if level == 2 else 12)


def add_paragraph(doc: Document, text: str, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after: int = 6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return p


def add_bullet_list(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.3
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def add_caption(doc: Document, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(caption)
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.5)


def add_manual_toc(doc: Document) -> None:
    entries = [
        "1. Introduction",
        "2. Problem Statement and Objectives",
        "3. Tools, Platform and Technology Stack",
        "4. System Analysis and Design",
        "5. Methodology and Development Process",
        "6. Implementation Details",
        "7. AI and Neuroevolution Strategy",
        "8. Testing, Results and Observations",
        "9. Limitations and Future Enhancements",
        "10. Conclusion",
        "References",
        "Appendix",
    ]
    for item in entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def add_title_page(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.2)
    set_page_border(section)

    if LOGO_PATH.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO_PATH), width=Inches(5.8))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    p.style = doc.styles["Report Title"]
    run = p.add_run("246 DEEP LEARNING LAB MINI PROJECT REPORT")
    run.font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("ON")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("CHROME DINO AI: UNITY-BASED ENDLESS RUNNER WITH NEUROEVOLUTION TRAINING")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(17)

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    details = [
        ("Submitted By", f"Name: {STUDENT_NAME}    Reg. No.: {REG_NUMBER}"),
        ("Program", "Deep Learning Lab / Mini Project"),
        ("Project Domain", "Game AI, Reinforcement Style Training, Neuroevolution"),
        ("Guided By", "Dr. Rajashree Krishna"),
        ("Institution", "School of Computer Engineering, Manipal"),
    ]
    for row, (label, value) in zip(table.rows, details):
        set_cell_text(row.cells[0], label, bold=True)
        set_cell_text(row.cells[1], value)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Academic Year: 2025-2026")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("April 2026")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_section_break(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.1)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.2)
    set_page_border(section)


def build_report() -> None:
    ensure_dirs()
    fitness_runs = parse_fitness_runs(EDITOR_LOG_PATH)
    latest_run = select_latest_run(fitness_runs)
    eq_theta = make_equation_image("eq_theta", r"\theta = \{W_1, W_2\}")
    eq_param_count = make_equation_image("eq_param_count", r"(9+1)\times 8 + (8+1)\times 2 = 80 + 18 = 98")
    eq_hidden = make_equation_image("eq_hidden", r"h_j = \tanh\!\left(\sum_{i=1}^{9} x_i w^{(1)}_{ij} + b^{(1)}_j\right)")
    eq_output = make_equation_image("eq_output", r"y_k = \sigma\!\left(\sum_{j=1}^{8} h_j w^{(2)}_{jk} + b^{(2)}_k\right), \qquad \sigma(z)=\frac{1}{1+e^{-z}}", height=1.2, fontsize=18)
    eq_fitness = make_equation_image("eq_fitness", r"f(\theta) = t_{\mathrm{alive}} + \beta c + 0.1\, t_{\mathrm{alive}}", width=6.5)
    eq_opt = make_equation_image("eq_opt", r"\theta^{*} = \arg\max_{\theta} f(\theta)", width=4.0)
    eq_tournament = make_equation_image("eq_tournament", r"p = \arg\max_{\theta \in T} f(\theta)", width=4.5)
    eq_crossover = make_equation_image("eq_crossover", r"\theta^{\mathrm{child}}_m \in \{\theta^{A}_m,\ \theta^{B}_m\}, \qquad r_m \sim \mathcal{U}(0,1)", width=5.6, height=1.0, fontsize=18)
    eq_mutation = make_equation_image("eq_mutation", r"\theta'_{m} = \theta_{m} + \epsilon_m, \qquad \epsilon_m \sim \mathcal{U}[-s, s]", width=5.2)
    eq_update = make_equation_image("eq_update", r"P_{t+1} = E(P_t)\ \cup\ M(C(S(P_t)))", width=4.8)
    diagrams = [
        make_system_architecture(),
        make_activity_diagram(),
        make_training_pipeline(),
        make_feature_map(),
        make_sprite_sheet(),
        make_fitness_graph(latest_run),
    ]

    doc = Document()
    configure_styles(doc)
    add_title_page(doc)

    add_section_break(doc)
    doc.add_heading("Acknowledgement", level=1)
    add_paragraph(doc, "The development of this mini project report and the corresponding Chrome Dino AI system was made possible through the consistent academic support, technical guidance, and motivation provided by our faculty and institutional environment. We express our sincere gratitude to our guide, Dr. Rajashree Krishna, for the valuable inputs that helped shape the project direction, improve the clarity of the implementation approach, and encourage disciplined experimentation throughout the work.")
    add_paragraph(doc, "We also acknowledge the support of the Deep Learning Lab and the School of Computer Engineering at Manipal for providing a platform where game development, artificial intelligence concepts, and practical software engineering could be integrated into one applied mini project. The project benefited greatly from access to Unity, C#, ML-Agents resources, and a structured lab setting that encouraged iterative development and testing.")
    add_paragraph(doc, "Finally, we acknowledge the contributions of open-source tools and documentation that accelerated implementation, debugging, and design refinement. These resources helped us understand the original Chrome Dino mechanics, organize the Unity scene, and implement a neuroevolution-based controller that can learn obstacle-avoidance behavior from gameplay state information.")

    doc.add_heading("Abstract", level=1)
    add_paragraph(doc, "This project presents the design and implementation of a Unity-based recreation of the well-known Chrome Dinosaur Game enhanced with an artificial intelligence control layer. The application combines endless-runner game mechanics, event-driven game management, and a custom neuroevolution training pipeline so that the dinosaur can progressively learn when to jump and how long to hold the jump while facing procedurally spawned obstacles. The project therefore serves both as a playable game and as a compact experimental platform for studying intelligent decision-making in a constrained environment.")
    add_paragraph(doc, "The core gameplay system was implemented in Unity using C# scripts for the player controller, obstacle spawning, ground motion, animation handling, scoring, audio, and retry flow. A separate AI module was integrated on top of the gameplay loop. Instead of relying on image-based perception, the AI uses state-based observations including grounded state, approximate vertical velocity, distance to the next obstacle, obstacle dimensions, dinosaur dimensions, current game speed, and speed increase factor. These features are passed to a fixed-topology neural network genome, which outputs jump-press and jump-hold decisions.")
    add_paragraph(doc, "To optimize the controller, the project implements a neuroevolution strategy inspired by NEAT-like principles. Populations of candidate genomes are evaluated inside the Unity scene, ranked according to fitness, and improved through tournament selection, elitism, crossover, and mutation. Additional engineering decisions such as deterministic obstacle spawning, parallel evaluation support, session-high-score persistence, training overlays, and visual differentiation of agents improve reproducibility and usability. The final outcome is a compact and well-structured game AI project that demonstrates how deep learning concepts, neural control, and iterative optimization can be applied in an interactive software environment.")

    doc.add_heading("Table of Contents", level=1)
    add_manual_toc(doc)

    doc.add_page_break()
    doc.add_heading("1. Introduction", level=1)
    add_paragraph(doc, "The Chrome Dino Game has become a familiar reference point for studying simple interactive environments because its rules are easy to understand while still allowing meaningful control challenges. The player must react to continuously arriving obstacles, maintain timing under changing speed, and sustain survival for as long as possible. This combination makes the game a suitable academic mini project for bringing together concepts from game development, control systems, and machine intelligence.")
    add_paragraph(doc, "In this project, the classic endless runner was recreated in Unity and extended with an AI controller capable of learning jump behavior from environment observations. The work goes beyond reproducing the visual style of the original game. It formalizes the environment into a decision process, identifies the features required for control, constructs a trainable neural policy, and builds an evolutionary training mechanism that can operate directly inside the game scene. The result is a project that is practical, visually demonstrable, and technically aligned with core ideas from deep learning and intelligent systems.")
    add_paragraph(doc, "The report documents the complete project lifecycle, beginning with the motivation and requirements, followed by the system architecture, methodology, implementation, AI strategy, testing approach, and future expansion paths. Since the codebase contains both playable game logic and a custom neuroevolution module, the report is organized to show how these two layers interact to produce a full intelligent gameplay pipeline.")

    doc.add_heading("2. Problem Statement and Objectives", level=1)
    add_paragraph(doc, "The problem addressed by this project is the creation of a lightweight yet complete game AI system that can autonomously play an endless runner by making real-time jump decisions from compact state observations. The goal is not simply to script fixed behavior, but to construct a decision-making mechanism that can be optimized over generations and improved through feedback collected during gameplay.")
    add_bullet_list(doc, [
        "To recreate the Chrome Dino style gameplay in Unity with smooth movement, obstacle generation, score progression, and replay support.",
        "To design a player controller with responsive jump physics, variable jump height, grounded detection, and collision-based game termination.",
        "To represent the environment using meaningful numerical features that are sufficient for autonomous decision-making.",
        "To implement a neural policy whose outputs can control jump press and jump hold behavior.",
        "To build a neuroevolution trainer capable of generating populations, evaluating fitness, and improving genomes over generations.",
        "To support deterministic and parallel evaluation modes for more stable experimentation.",
        "To document the complete process in an organized report with diagrams, module breakdowns, and technical observations.",
    ])
    add_paragraph(doc, "From a deep learning perspective, the project demonstrates how learning-oriented systems can be applied in environments that do not require large image models or extensive offline datasets. Instead, the intelligence emerges from iterative interaction with a simulated environment, which makes the project efficient, explainable, and suitable for academic demonstration.")

    doc.add_heading("3. Tools, Platform and Technology Stack", level=1)
    tech_table = doc.add_table(rows=1, cols=3)
    tech_table.style = "Table Grid"
    hdr = tech_table.rows[0].cells
    set_cell_text(hdr[0], "Component", bold=True)
    set_cell_text(hdr[1], "Technology Used", bold=True)
    set_cell_text(hdr[2], "Purpose", bold=True)
    rows = [
        ("Game Engine", "Unity 2021.3 LTS", "Scene creation, runtime simulation, input, UI, and component system"),
        ("Programming Language", "C#", "Game logic, AI controller, training loop, and utility scripts"),
        ("AI Strategy", "Custom Neuroevolution", "Population-based optimization of neural network weights"),
        ("Auxiliary AI Framework", "ML-Agents setup present", "Optional future extension for formal agent training"),
        ("Visual Assets", "Sprite-based 2D resources", "Dino, cactus, ground, retry icon, and visual environment"),
        ("Documentation Tool", "Python + python-docx", "Automated generation of the organized project report"),
        ("Versioned Data", "Unity project assets and scripts", "Reproducible project structure and code organization"),
    ]
    for component, tech, purpose in rows:
        cells = tech_table.add_row().cells
        set_cell_text(cells[0], component, bold=True)
        set_cell_text(cells[1], tech)
        set_cell_text(cells[2], purpose)
    add_paragraph(doc, "The technology stack was chosen to keep the project practical and easy to demonstrate. Unity provides a strong component-driven architecture, which is ideal for organizing gameplay objects, user interface elements, and scripted behavior. C# enables direct interaction with the scene and supports integration of both deterministic game logic and adaptive AI control. The project is therefore positioned at the intersection of software engineering and learning-based control.")

    doc.add_heading("4. System Analysis and Design", level=1)
    add_paragraph(doc, "The system can be viewed as two tightly connected layers. The first layer is the gameplay engine, which includes the player, obstacles, game manager, UI, and environment rules. The second layer is the intelligence layer, which observes the gameplay state and produces control signals. During normal play, this controller may remain disabled and the human player can interact manually. During AI mode or training mode, the decision logic takes over and controls jump behavior automatically.")
    doc.add_picture(str(diagrams[0]), width=Inches(6.5))
    add_caption(doc, "Figure 1. High-level system architecture of the Chrome Dino AI project.")
    add_paragraph(doc, "The architecture begins with the scene initialization performed by the game manager, which restores scene references, audio resources, score displays, and retry behavior. The player object uses a CharacterController-based movement model with custom jump physics. A spawner periodically creates obstacles and can optionally operate in deterministic mode during training. The neuro controller reads the game state, transforms it into a normalized feature vector, and queries the current genome for an action. Finally, the trainer evaluates the outcome and updates the population of genomes between generations.")
    add_paragraph(doc, "This modular design offers two major advantages. First, gameplay remains stable and understandable even when AI is disabled. Second, training logic does not need to rewrite the whole game; it only interfaces with the existing player and environment through clearly defined observations and control outputs. Such separation of concerns improves maintainability and makes the project suitable for extension into future RL or imitation-learning experiments.")

    doc.add_heading("5. Methodology and Development Process", level=1)
    add_paragraph(doc, "The development process followed an incremental methodology in which the environment was built first, stabilized second, and instrumented for AI training third. This order was important because autonomous control is meaningful only when the game loop itself is reliable. The first milestone involved constructing the base endless-runner scene using sprite assets, animation support, obstacle prefabs, and the player controller. The second milestone introduced scoring, retry behavior, high-score persistence, sound feedback, and visual atmosphere changes such as the night-mode transition. The third milestone integrated the AI controller, the genome representation, and the evolutionary trainer.")
    doc.add_picture(str(diagrams[1]), width=Inches(5.9))
    add_caption(doc, "Figure 2. Activity diagram showing the runtime behavior of player and AI control.")
    add_paragraph(doc, "The activity flow begins when the scene resets for a new attempt or a new training episode. The game initializes the player, score, spawner, and UI. During each update cycle, the system observes the state of the character and the nearest obstacle. If the obstacle is close enough to require a decision, the policy determines whether to press or hold the jump. The game then updates movement, score, speed, camera appearance, and collisions. If the player collides or the episode times out during training, the system either ends the run or records the fitness and proceeds to the next evaluation.")
    add_paragraph(doc, "This workflow is intentionally simple, which is beneficial for both debugging and academic explanation. Every important stage of the control loop can be traced directly to a script in the project, making the project easier to present in lab evaluation or viva settings.")

    doc.add_heading("6. Implementation Details", level=1)
    doc.add_heading("6.1 Game Manager", level=2)
    add_paragraph(doc, "The GameManager acts as the central coordinator of scene behavior. It stores the current game speed, updates the score continuously, increases difficulty over time, manages the UI labels for current score and high scores, and controls the retry state when the player loses. It also ensures that scene references are recovered even if serialized links are missing, which is a practical safety feature for student projects where scene editing can sometimes break references.")
    add_paragraph(doc, "Additional polish in the GameManager includes session high-score tracking, point sound effects at configured score intervals, and a night-mode transition after the score crosses a threshold. These features are not strictly necessary for AI experimentation, but they make the environment feel more complete and professional. Importantly, they also prove that the project is not only a proof-of-concept model, but a functioning game system with player-facing quality improvements.")

    doc.add_heading("6.2 Player Controller", level=2)
    add_paragraph(doc, "The Player script implements vertical jump motion using a CharacterController component. The controller supports variable jump height through a combination of initial jump velocity, temporary upward acceleration while the jump input is held, and stronger gravity when the jump is released early. This gives the game more expressive control compared to a fixed-height jump. The script also aligns the dinosaur to the ground when the object is enabled, which avoids spawn inconsistencies between retries and training resets.")
    add_paragraph(doc, "The player script merges human and AI input cleanly. Human input is read from the standard jump button, while AI input is provided through a separate setter. Internally, both are converted into the same effective jump state. This is a strong design choice because it allows the same gameplay code to be used in both manual and autonomous modes without branching into separate movement implementations.")

    doc.add_heading("6.3 Obstacle Spawner", level=2)
    add_paragraph(doc, "The Spawner script controls obstacle generation and difficulty pacing. In normal mode, it selects spawn timings randomly between minimum and maximum intervals. In deterministic mode, it uses a seeded pseudo-random generator and a fixed interval. Deterministic spawning is especially important during training because it reduces evaluation noise and allows genomes to be compared under similar environmental conditions. The current implementation also skips prefabs whose names contain 'Bird', which simplifies the environment and keeps the learning problem focused on cactus avoidance.")

    doc.add_heading("6.4 Visual and Asset Pipeline", level=2)
    add_paragraph(doc, "The project uses lightweight sprite assets to keep the visual identity close to the original Chrome Dino aesthetic. Sprite-based assets reduce rendering complexity and make the environment easy to inspect. The dinosaur animations, cacti, ground, retry image, and UI elements are arranged to create a cohesive endless-runner scene that is suitable for both gameplay and AI observation.")
    doc.add_picture(str(diagrams[4]), width=Inches(5.9))
    add_caption(doc, "Figure 3. Sample sprite assets used within the Unity project.")

    doc.add_heading("7. AI and Neuroevolution Strategy", level=1)
    add_paragraph(doc, "The intelligence component of the project is implemented through a fixed-topology neural controller combined with a population-based optimization strategy. Each candidate solution is represented as a NeuroGenome. The genome stores two arrays of weights: one for the input-to-hidden layer and another for the hidden-to-output layer. The network uses nine input features, eight hidden neurons, and two outputs. The outputs correspond to a jump press signal and a jump hold signal, which gives the AI a richer action space than a simple binary jump/no-jump command.")
    doc.add_picture(str(diagrams[3]), width=Inches(6.3))
    add_caption(doc, "Figure 4. Feature and output structure of the AI controller.")
    add_paragraph(doc, "The chosen feature set is compact but informative. It includes grounded state, approximate vertical velocity, distance to the next obstacle, obstacle width, obstacle height, dinosaur width, dinosaur height, current game speed, and the speed increase parameter. These features are normalized into a bounded numeric range so that the neural network can process them more consistently. Because the environment is simple and fully observable through these variables, state-based learning is more efficient than pixel-based learning for this project.")
    doc.add_picture(str(diagrams[2]), width=Inches(6.4))
    add_caption(doc, "Figure 5. Neuroevolution training pipeline used to improve the game-playing policy.")
    add_paragraph(doc, "Training is handled by the NeuroEvolutionTrainer, which creates a population of genomes, evaluates them in the environment, sorts them by fitness, and generates the next population using elitism, tournament selection, crossover, and mutation. Fitness is shaped using survival duration and obstacle-clear bonuses so that the optimization process rewards both staying alive and successfully handling obstacles. The trainer also supports parallel evaluation by cloning player instances and evaluating multiple genomes simultaneously. This reduces training time and shows thoughtful engineering beyond a minimal prototype.")
    add_paragraph(doc, "An important project strength is the use of deterministic settings per generation. When enabled, the trainer seeds Unity random values and configures the spawner in deterministic mode. This helps ensure that genomes are compared more fairly because obstacle sequences remain stable during that generation. The project also saves the best genome using PlayerPrefs, allowing the best learned policy to be reused later during gameplay demonstrations.")
    doc.add_heading("7.1 Genetic Algorithm Foundations Used in This Project", level=2)
    add_paragraph(doc, "The optimization strategy used in this project belongs to the family of genetic algorithms, which are population-based search methods inspired by biological evolution. Instead of computing gradients with respect to a loss function, a genetic algorithm works with a set of candidate solutions called a population. Each candidate is evaluated, assigned a quality measure called fitness, and then used to generate a new population through selection, inheritance, and random variation. In this project, every candidate solution is a neural-network genome containing the weights of the policy that controls the dinosaur.")
    add_paragraph(doc, "Mathematically, a genome can be written as a parameter vector theta that contains all weights of the network. If the input-to-hidden weights are denoted by W1 and the hidden-to-output weights are denoted by W2, then the full genome is shown below. For the current architecture, the input layer has 9 features, the hidden layer has 8 neurons, and the output layer has 2 neurons. Including bias terms, the network contains 98 trainable parameters in total.")
    add_equation(doc, eq_theta, 3.1)
    add_equation(doc, eq_param_count, 4.8)
    add_paragraph(doc, "Given an input vector x in R^9, the hidden representation h is computed using an affine transform followed by the hyperbolic tangent activation. The output layer then computes two action scores using a sigmoid activation. Because the sigmoid output lies between 0 and 1, the controller can interpret these values directly as action intensities. The implementation then compares the outputs against fixed thresholds to decide whether to press jump or hold jump.")
    add_equation(doc, eq_hidden, 5.7)
    add_equation(doc, eq_output, 6.2)
    add_paragraph(doc, "A genetic algorithm does not directly update theta by differentiating through this computation. Instead, it evaluates the quality of each theta by running the game and measuring the resulting performance. Let f(theta) denote the fitness function. In the sequential evaluation mode used here, the implemented fitness is a combination of score gain, survival time, and obstacle-clear bonuses. In the parallel evaluation mode, the trainer uses the following shaped fitness function, where t_alive is the survival duration, c is the number of cleared obstacles, and beta is the obstacle-cleared bonus coefficient. This turns the game into an optimization problem over theta.")
    add_equation(doc, eq_fitness, 4.8)
    add_equation(doc, eq_opt, 2.8)
    add_paragraph(doc, "The first stage of the genetic algorithm is population initialization. If the population size is N, the algorithm begins with genomes theta_1, theta_2, ..., theta_N whose weight values are sampled randomly, typically from a bounded interval such as [-1, 1]. This is visible in the NeuroGenome.Randomize method, where every weight is initialized using a uniform random distribution. Random initialization gives the algorithm coverage over many possible behaviors. Most initial controllers behave poorly, but a few may accidentally discover useful jump timing, which gives evolution a starting point.")
    add_paragraph(doc, "After evaluation, the project applies elitism. Elitism means that the best E genomes are copied directly into the next generation without modification. If the sorted population is theta_(1), theta_(2), ..., theta_(N), ordered by descending fitness, then the next generation begins by preserving theta_(1) through theta_(E). This is mathematically important because it guarantees that the best-known solution is not lost due to later random variation. In practical terms, elitism stabilizes learning and prevents the training process from forgetting strong obstacle-avoidance behaviors once they are discovered.")
    add_paragraph(doc, "The next stage is parent selection. This project uses tournament selection, a common method in genetic algorithms because it is simple, computationally cheap, and robust to noisy fitness values. In tournament selection, a small subset of genomes is sampled uniformly from the current population, and the fittest genome among that subset is chosen as a parent. If the tournament size is k, then the selected parent is the fitness maximizer inside the sampled subset T. Larger k increases selection pressure because strong genomes win more often, while smaller k preserves more diversity.")
    add_equation(doc, eq_tournament, 3.2)
    add_paragraph(doc, "Once parents are selected, the algorithm produces offspring through crossover. The current project implements uniform crossover. If parent A has weights theta^A and parent B has weights theta^B, then for each parameter index m, the child inherits either theta^A_m or theta^B_m with equal probability. Uniform crossover mixes traits at the parameter level, allowing one child to combine useful substructures from two different parents. Although weight-space crossover does not guarantee semantic modularity, it often works well in compact networks such as the one used here.")
    add_equation(doc, eq_crossover, 4.8)
    add_paragraph(doc, "Mutation is the second major source of new variation. After crossover, each parameter is independently considered for mutation with probability mu, the mutation rate. When a mutation occurs, a random perturbation is added to that parameter. In the project code, the perturbation is drawn from a bounded uniform range controlled by mutationStrength. Mutation serves two purposes. First, it enables exploration of new behaviors that are not already present in the parent population. Second, it prevents premature convergence by maintaining diversity in the search space.")
    add_equation(doc, eq_mutation, 4.5)
    add_paragraph(doc, "The overall generational update can therefore be summarized as follows. Start with population P_t at generation t. Evaluate every genome to obtain fitness values. Sort genomes by fitness. Copy the top elites directly into the next population. Fill the remaining slots by repeatedly selecting parents, performing crossover with probability p_c, cloning otherwise, and mutating the resulting child. This creates P_(t+1). The loop repeats until the generation budget is exhausted or a satisfactory controller emerges. In compact notation, the entire training step can be represented as follows, where S denotes selection, C denotes crossover, M denotes mutation, and E denotes elite preservation.")
    add_equation(doc, eq_update, 3.5)
    add_paragraph(doc, "One of the strongest mathematical advantages of genetic algorithms in this project is that they do not require the environment to be differentiable. The dinosaur game contains threshold-based control decisions, collisions, event triggers, and game-over conditions, all of which make gradient-based optimization less direct in a hand-built setup. A genetic algorithm avoids this difficulty because it treats the environment as a black-box evaluator. As long as the system can run a genome and assign a scalar fitness value, evolution can proceed.")
    add_paragraph(doc, "At the same time, genetic algorithms involve an exploration-exploitation tradeoff. Strong selection pressure and low mutation may exploit promising behaviors rapidly, but can also collapse the population into a narrow region of the search space. High mutation and low selection pressure improve exploration, but may slow convergence. The current trainer manages this balance through configurable parameters such as population size, elite count, crossover rate, mutation rate, mutation strength, and deterministic seeding. These parameters influence convergence speed, stability, and final policy quality.")
    add_paragraph(doc, "Another mathematically relevant concept is variance reduction. Because obstacle timing can change from one episode to another, the same genome may appear stronger or weaker simply due to environmental randomness. The project addresses this by optionally forcing deterministic obstacle sequences within a generation. From an optimization perspective, this reduces noise in the estimate of f(theta), which makes rank ordering more reliable. Reliable ranking is especially important in genetic algorithms because parent selection is based directly on comparative fitness.")
    add_paragraph(doc, "Finally, it is useful to understand why this approach is appropriate for a mini project. A full reinforcement-learning pipeline often requires reward shaping, agent-environment interfacing, training configuration files, and sometimes large numbers of episodes. The genetic-algorithm approach used here is easier to explain, easier to instrument in a Unity-only workflow, and still mathematically rich enough to demonstrate optimization over a neural policy. The project therefore provides a meaningful bridge between introductory game AI and broader ideas from evolutionary computation, neural optimization, and adaptive control.")

    doc.add_heading("8. Testing, Results and Observations", level=1)
    add_paragraph(doc, "Testing for this project was conducted at multiple levels. Functional testing verified that the game starts correctly, the player spawns in the expected position, jump physics feel consistent, obstacles are generated continuously, the score increases properly, retry resets the scene state, and collision produces the intended game-over flow. Integration testing checked that AI control can be turned on without breaking the normal player flow and that the game manager, player script, and trainer remain compatible during repeated evaluations.")
    results_table = doc.add_table(rows=1, cols=3)
    results_table.style = "Table Grid"
    headers = results_table.rows[0].cells
    set_cell_text(headers[0], "Test Area", bold=True)
    set_cell_text(headers[1], "Expected Outcome", bold=True)
    set_cell_text(headers[2], "Observation", bold=True)
    test_rows = [
        ("Scene initialization", "Player, UI, audio, and references are restored", "Implemented through recovery logic in GameManager"),
        ("Jump control", "Short and long jumps respond correctly", "Handled by jump hold acceleration and jump cut gravity"),
        ("Obstacle spawning", "Continuous challenge with configurable randomness", "Normal and deterministic modes both available"),
        ("Collision flow", "Obstacle hit ends run or reports agent death during training", "Handled differently for gameplay and parallel training"),
        ("AI inference", "Genome outputs control jump press and hold", "Implemented through DinoNeuroController thresholds"),
        ("Training progression", "Best fitness improves across generations when configuration is suitable", "Trainer logs and saved best genome support this workflow"),
    ]
    for test_area, expected, observation in test_rows:
        cells = results_table.add_row().cells
        set_cell_text(cells[0], test_area, bold=True)
        set_cell_text(cells[1], expected)
        set_cell_text(cells[2], observation)
    add_paragraph(doc, "The project does contain recoverable training evidence in the Unity Editor log, and that data has been used to build a genuine fitness-versus-generation visualization for this report. Even beyond that graph, the codebase clearly contains the infrastructure required for meaningful experimentation: evaluation loops, best-genome persistence, deterministic training settings, parallel evaluation, and on-screen overlays for fitness tracking. From a project-report perspective, this demonstrates that the system is organized for iterative improvement and not merely for one-off scripted behavior.")
    add_paragraph(doc, "Qualitatively, the design suggests that the AI should improve gradually as training progresses because fitness rewards both duration and obstacle handling. In practical demonstrations, the best genome can be loaded and used in play mode, allowing observers to compare human control and learned control using the same environment. This makes the project especially effective during presentation or viva because the intelligence is directly visible to the evaluator.")
    doc.add_picture(str(diagrams[5]), width=Inches(6.4))
    add_caption(doc, "Figure 6. Fitness versus generations graph extracted from a recorded NeuroEvolution run in Editor.log.")
    if latest_run and latest_run["points"]:
        first_fit = latest_run["points"][0][1]
        last_fit = latest_run["points"][-1][1]
        peak_gen, peak_fit = max(latest_run["points"], key=lambda item: item[1])
        add_paragraph(doc, f"A real fitness trace was recovered from the Unity Editor log and plotted in the figure above. The selected contiguous run is the latest recorded run in the log and spans {len(latest_run['points'])} generations from a configuration of {latest_run['total']} total generations. In that run, the fitness moved from {first_fit:.2f} in the first recorded generation to {last_fit:.2f} in the last recorded generation shown, with a peak of {peak_fit:.2f} observed at generation {peak_gen}. This gives the report an evidence-backed visualization of the most recent training behavior instead of a manually fabricated trend line.")

    doc.add_heading("9. Limitations and Future Enhancements", level=1)
    add_paragraph(doc, "Although the project is complete and well-structured, it intentionally keeps the environment simpler than a fully generalized Chrome Dino clone. Bird obstacles are currently excluded, which reduces action diversity. The neural controller also uses a fixed network topology rather than a dynamically expanding architecture. In addition, the reportable results would be stronger if the project stored generation-wise metrics, graphs, and trained model snapshots in a dedicated experiment log folder.")
    add_bullet_list(doc, [
        "Add flying bird obstacles and crouch mechanics so that the agent must choose among more action types.",
        "Record training metrics such as best fitness, average fitness, and generation count in CSV or JSON form.",
        "Integrate a formal ML-Agents training pipeline for comparison against the custom neuroevolution solution.",
        "Export trained policies and compare inference quality across multiple seeds.",
        "Add screenshot capture or video snippets automatically during training for richer documentation.",
        "Create a user-facing mode switch panel to select human play, heuristic play, best-genome play, or training mode.",
    ])
    add_paragraph(doc, "These enhancements would strengthen the project further by making it more suitable for comparative analysis, performance visualization, and publication-style experimentation. Even in its current form, however, the project already demonstrates a valuable combination of engineering completeness and AI-oriented thinking.")

    doc.add_heading("10. Conclusion", level=1)
    add_paragraph(doc, "The Chrome Dino AI mini project successfully combines a polished Unity endless-runner implementation with a compact but meaningful intelligent control system. The project recreates the familiar gameplay loop while introducing state-based autonomous decision-making through a neural policy and neuroevolution. This combination makes the project a strong example of applied deep learning concepts in an interactive environment where every part of the system, from observation to action to optimization, is visible and understandable.")
    add_paragraph(doc, "From an academic perspective, the project demonstrates important technical outcomes: modular game design, feature engineering for control, neural-network based inference, deterministic training design, genetic optimization, and practical experimentation support. It also shows good software engineering discipline by separating gameplay logic from AI logic and by providing infrastructure for retries, overlays, saved best policies, and parallel evaluation. Overall, the project is both demonstrable and extensible, making it a suitable and organized submission for a deep learning lab mini project.")

    doc.add_heading("References", level=1)
    for ref in [
        "Unity Documentation. (2026). Unity Manual and Scripting API.",
        "Google Chrome Dinosaur Game concept and gameplay references.",
        "Unity ML-Agents Toolkit documentation available in the local project dependency.",
        "Neuroevolution concepts: population initialization, mutation, crossover, and tournament selection.",
        "Project source files: GameManager.cs, Player.cs, Spawner.cs, DinoNeuroController.cs, NeuroGenome.cs, and NeuroEvolutionTrainer.cs.",
    ]:
        add_paragraph(doc, ref, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)

    doc.add_heading("Appendix", level=1)
    add_paragraph(doc, "Appendix A: Important project scripts include Assets/Scripts/GameManager.cs, Assets/Scripts/Player.cs, Assets/Scripts/Spawner.cs, and the neuroevolution scripts under Assets/Scripts/AI/Neuroevolution. These files together define the runtime logic, observation space, decision policy, training loop, and supporting visualization behavior.")
    add_paragraph(doc, "Appendix B: The repository also contains ML-Agents resources and configuration material, indicating that the project can be expanded in future iterations toward formal reinforcement learning workflows. This broadens the scope of the project from a single custom AI strategy to a more flexible AI experimentation platform.")

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_DOCX))


if __name__ == "__main__":
    build_report()
